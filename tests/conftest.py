import io
import pytest
from docx import Document
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

@pytest.fixture
def sample_pdf_bytes():
    """Create a minimal valid PDF with selectable text."""
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.drawString(100, 750, "John Doe")
    p.drawString(100, 730, "john.doe@example.com | +1-202-555-0123")
    p.drawString(100, 710, "Summary: Experienced Software Engineer with Python and React skills.")
    p.drawString(100, 690, "EXPERIENCE")
    p.drawString(100, 670, "Senior Developer - Acme Corp (2020 - Present)")
    p.drawString(100, 650, "- Developed scalable backend services using FastAPI and PostgreSQL.")
    p.drawString(100, 630, "- Improved API response times by 40% using Redis caching.")
    p.drawString(100, 610, "SKILLS")
    p.drawString(100, 590, "Python, FastAPI, Docker, SQL, JavaScript, React")
    p.drawString(100, 570, "PROJECTS")
    p.drawString(100, 550, "E-commerce API: Built full-stack store using Python and Docker.")
    p.drawString(100, 530, "EDUCATION")
    p.drawString(100, 510, "B.S. in Computer Science - Tech University (2016 - 2020)")
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.getvalue()

@pytest.fixture
def sample_docx_bytes():
    """Create a minimal valid DOCX with structured text."""
    doc = Document()
    doc.add_heading("Jane Smith", level=1)
    doc.add_paragraph("jane.smith@example.com | +1-202-555-9876")
    doc.add_heading("Summary", level=2)
    doc.add_paragraph("Full-Stack Developer passionate about cloud architecture.")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Software Engineer - CloudTech (2021 - Present)")
    doc.add_paragraph("- Engineered distributed microservices with Python and Docker.")
    doc.add_paragraph("- Optimized database queries saving 25% CPU utilization.")
    doc.add_heading("Skills", level=2)
    doc.add_paragraph("Python, Django, AWS, Kubernetes, PostgreSQL")
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("Analytics Dashboard: Visualized metrics using Python and PostgreSQL.")
    doc.add_heading("Education", level=2)
    doc.add_paragraph("B.Tech in Information Technology - Global Institute (2017 - 2021)")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@pytest.fixture
def sample_resume_text():
    return """
    John Doe
    john.doe@example.com | +1-555-0199 | linkedin.com/in/johndoe | github.com/johndoe

    PROFESSIONAL SUMMARY
    Full-Stack Software Engineer with 4+ years of experience designing and deploying high-performance web applications using Python, FastAPI, and React.

    EXPERIENCE
    Senior Backend Developer — TechNova Inc (Jan 2022 – Present)
    • Developed REST APIs with FastAPI serving 50,000 daily active users.
    • Optimized database performance in PostgreSQL, reducing query latency by 45%.
    • Automated CI/CD deployment pipelines using Docker and GitHub Actions.

    Software Engineer — StartupX (Jun 2020 – Dec 2021)
    • Built real-time analytics dashboard with React and Python.
    • Implemented Redis caching layer improving throughput by 35%.

    PROJECTS
    • AI ATS Resume Analyzer — Built an intelligent resume evaluator using Python, FastAPI, and Groq LLMs. Handled 500+ daily analyses.
    • Cloud Task Manager — Engineered a collaborative task management app with React, Node.js, and MongoDB.

    TECHNICAL SKILLS
    Languages & Frameworks: Python, JavaScript, TypeScript, FastAPI, React, Django
    Tools & Databases: PostgreSQL, MongoDB, Redis, Docker, Git, AWS

    EDUCATION
    B.S. in Computer Science — University of Technology (2016 – 2020)
    CGPA: 3.8/4.0
    """

@pytest.fixture
def sample_parsed_resume():
    return {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "+1-555-0199",
        "linkedin": "linkedin.com/in/johndoe",
        "github": "github.com/johndoe",
        "professional_summary": "Full-Stack Software Engineer with 4+ years of experience designing web applications.",
        "skills": ["Python", "FastAPI", "React", "PostgreSQL", "Docker", "Redis", "Git", "AWS"],
        "experience": [
            {
                "job_title": "Senior Backend Developer",
                "company": "TechNova Inc",
                "start_date": "2022-01",
                "end_date": "Present",
                "duration_months": 24,
                "description": "Developed REST APIs with FastAPI serving 50K users. Optimized PostgreSQL queries by 45%."
            }
        ],
        "education": [
            {
                "degree": "B.S. in Computer Science",
                "institution": "University of Technology",
                "year": "2020"
            }
        ],
        "certifications": ["AWS Certified Solutions Architect"],
        "projects": [
            {
                "title": "AI ATS Resume Analyzer",
                "description": "Built an intelligent resume evaluator using Python, FastAPI, and Groq LLMs.",
                "technologies": ["Python", "FastAPI", "Docker"]
            }
        ],
        "action_verbs": ["Developed", "Optimized", "Automated", "Built", "Implemented", "Engineered"],
        "keywords": ["Python", "FastAPI", "React", "PostgreSQL", "Docker", "REST API", "Microservices"]
    }
